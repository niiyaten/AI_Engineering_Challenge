from __future__ import annotations

import argparse
import csv
import json
import re
import subprocess
import unicodedata
from collections import Counter, defaultdict
from decimal import Decimal, ROUND_HALF_UP
from pathlib import Path
from typing import Any

from docx import Document


CALCULATION_PATTERNS = (
    "single_table_filter", "single_table_aggregation", "multi_condition_aggregation",
    "multi_step_calculation", "difference_calculation", "ratio_or_percentage",
    "hypothetical_calculation", "coefficient_prediction", "pivot_lookup",
    "schedule_effort_aggregation", "id_count_or_nunique", "cross_file_join",
    "cross_file_aggregation", "date_or_duration_calculation", "ranking_or_argmin_argmax",
    "chart_source_calculation", "not_actually_calculation", "unknown_calculation",
)


def normalize(value: Any) -> str:
    return re.sub(r"\s+", " ", unicodedata.normalize("NFKC", str(value or ""))).strip()


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def read_jsonl(path: Path) -> list[dict[str, Any]]:
    if not path.exists():
        return []
    return [json.loads(line) for line in path.read_text(encoding="utf-8").splitlines() if line.strip()]


def write_csv(path: Path, rows: list[dict[str, Any]], columns: list[str] | None = None) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    fields = columns or (list(rows[0]) if rows else ["status"])
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fields, extrasaction="ignore")
        writer.writeheader(); writer.writerows(rows)


def write_jsonl(path: Path, rows: list[dict[str, Any]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text("\n".join(json.dumps(row, ensure_ascii=False) for row in rows) + ("\n" if rows else ""), encoding="utf-8")


def classify_pattern(question: str) -> tuple[str, list[str]]:
    """質問の演算語と情報源構成だけからCalculationパターンを分類する。"""
    q = normalize(question)
    secondary: list[str] = []
    if any(term in q for term in ("残余リスク", "スコープ対象外")) and not any(term in q for term in ("計算", "件数", "いくつ")):
        return "not_actually_calculation", []
    if any(term in q for term in ("回帰係数", "係数を使", "係数を用い")):
        primary = "coefficient_prediction"
    elif any(term in q for term in ("仮に", "だった場合", "高く", "少なかった場合")):
        primary = "hypothetical_calculation"
    elif "Pivot" in q or "PivotTable" in q:
        primary = "pivot_lookup"
    elif any(term in q for term in ("差額", "改善幅", "差を", "減額", "変動")):
        primary = "difference_calculation"
    elif any(term in q for term in ("割合", "何%", "何倍", "上昇率")):
        primary = "ratio_or_percentage"
    elif any(term in q for term in ("タスクIDはいくつ", "IDは合計でいくつ", "全部で何人", "いくつありますか")):
        primary = "id_count_or_nunique"
    elif "工数" in q and any(term in q for term in ("合計", "1タスク", "担当")):
        primary = "schedule_effort_aggregation"
    elif any(term in q for term in ("最も高い", "最も低い", "最も近い", "4番目")):
        primary = "ranking_or_argmin_argmax"
    elif "グラフ" in q or "ヒストグラム" in q or "可視化" in q:
        primary = "chart_source_calculation"
    elif sum(term in q for term in ("かつ", "の中で", "に該当", "より大きい", "未満")) >= 2:
        primary = "multi_condition_aggregation"
    elif any(term in q for term in ("平均", "合計", "件数")):
        primary = "single_table_aggregation"
    else:
        primary = "unknown_calculation"

    if any(term in q for term in ("と、", "と最終", "中間報告", "全案件", "各案件", "提案時", "最終報告", "社内管理")):
        secondary.append("cross_file_join")
    if any(term in q for term in ("平均を上回る", "最も近い", "閾値を設定", "割った")):
        secondary.append("multi_step_calculation")
    if any(term in q for term in ("小数第", "四捨五入", "切り上げ")):
        secondary.append("rounding")
    return primary, secondary


def required_fields(question: str) -> tuple[list[str], list[str], list[str], list[str]]:
    q = normalize(question)
    columns = re.findall(r"\b[A-Za-z][A-Za-z0-9_]*\b", q)
    columns = [item for item in columns if item.lower() not in {"csv", "xlsx", "json", "pivot", "sheet", "macro", "f1", "id"}]
    conditions = re.findall(r"[A-Za-z][A-Za-z0-9_]*\s*(?:=|が)\s*[^、。]+", q)
    units = [unit for unit in ("円", "%", "時間", "日", "歳", "件", "人") if unit in q]
    rounding = re.findall(r"小数第\s*\d+\s*位|四捨五入|切り上げ|整数値", q)
    return list(dict.fromkeys(columns)), conditions, units, rounding


def source_shape(question: str) -> tuple[str, str]:
    q = normalize(question)
    if any(term in q for term in ("全案件", "各案件", "完了案件")):
        return "all_matching", "cross_project"
    if any(term in q for term in ("中間報告", "提案時", "最終報告時点", "metrics.json", "社内管理")) and any(term in q for term in ("差", "改善幅", "合計", "最も")):
        return "multiple", "aggregate_sources"
    return "single", "same_project"


def docx_text(path: Path) -> str:
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        parts.extend(" | ".join(cell.text for cell in row.cells) for row in table.rows)
    return "\n".join(parts)


def metric_value(text: str, metric: str, phase: str = "") -> Decimal | None:
    aliases = [r"Macro\s*F1", r"F1\s*\(macro\)"] if metric == "f1_macro" else [r"Accuracy"]
    matches: list[tuple[int, Decimal]] = []
    for alias in aliases:
        for found in re.finditer(alias + r"\s*[:=：]?\s*([01](?:\.\d+)?)", text, re.I):
            context = text[max(0, found.start() - 180): found.end() + 80]
            score = (5 if phase and phase in context else 0) + (3 if any(word in context for word in ("最良", "公開可能", "サマリー", "共有値")) else 0)
            matches.append((score, Decimal(found.group(1))))
    if not matches:
        return None
    best_score = max(score for score, _ in matches)
    values = {value for score, value in matches if score == best_score}
    return next(iter(values)) if len(values) == 1 else None


def generate_synthetic(root: Path) -> tuple[list[dict[str, Any]], Path]:
    base = root / "evaluation/synthetic/calculation/cross_source_metric_difference"
    cases = [
        ("positive_basic", True, "0.700000", "0.745000", "0.045000", "basic two-source difference"),
        ("positive_source_order", True, "0.610000", "0.650000", "0.040000", "unrelated file and source order change"),
        ("positive_row_order", True, "0.660000", "0.710000", "0.050000", "nested field order change"),
        ("positive_multiple_conditions", True, "0.720000", "0.780000", "0.060000", "project, phase, and metric conditions"),
        ("positive_numeric_string", True, "0.800000", "0.812346", "0.012346", "numeric strings"),
        ("positive_missing_optional", True, "0.640000", "0.675000", "0.035000", "optional fields missing"),
        ("positive_extra_fields", True, "0.500000", "0.575000", "0.075000", "extra fields"),
        ("positive_unrelated_file", True, "0.690000", "0.715000", "0.025000", "unrelated file ignored"),
        ("negative_missing_metric", False, "", "0.700000", "", "required metric missing"),
        ("negative_missing_source", False, "0.700000", "", "", "required source missing"),
        ("negative_ambiguous_source", False, "0.700000|0.710000", "0.730000", "", "conflicting source values"),
        ("negative_unit_mismatch", False, "72%", "0.75", "", "unit mismatch"),
        ("negative_empty_target_value", False, "0.700000", "", "", "condition exists but target value is empty"),
        ("negative_duplicate_join_key", False, "0.700000|0.705000", "0.730000", "", "duplicate source role key"),
        ("negative_numeric_id", False, "", "", "", "numeric ID must not be used as metric"),
        ("negative_operation_mismatch", False, "0.700000", "0.750000", "", "question requests ratio instead of difference"),
        ("negative_unrelated_project", False, "0.700000", "0.750000", "", "source relation not verified"),
    ]
    rows: list[dict[str, Any]] = []
    for case_id, positive, before, after, expected, note in cases:
        case_dir = base / case_id; case_dir.mkdir(parents=True, exist_ok=True)
        before_value: Any = before.split("|") if "|" in before else before
        interim_payload = {"phase": "interim", "f1_macro": before_value, "row_id": 1001}
        final_payload = {"phase": "final", "f1_macro": after, "unused": 999, "row_id": 1002}
        if case_id == "positive_row_order":
            interim_payload = {"metadata": {"note": "order changed"}, "f1_macro": before, "phase": "interim"}
            final_payload = {"f1_macro": after, "phase": "final", "metadata": {"note": "order changed"}}
        if case_id == "positive_missing_optional":
            interim_payload = {"phase": "interim", "f1_macro": before}
            final_payload = {"phase": "final", "f1_macro": after}
        if case_id == "negative_numeric_id":
            interim_payload = {"phase": "interim", "id": 700000}
            final_payload = {"phase": "final", "id": 750000}
        (case_dir / "interim.json").write_text(json.dumps(interim_payload, ensure_ascii=False, indent=2), encoding="utf-8")
        (case_dir / "final.json").write_text(json.dumps(final_payload, ensure_ascii=False, indent=2), encoding="utf-8")
        (case_dir / "unrelated.json").write_text(json.dumps({"project": "other", "f1_macro": 0.99}, ensure_ascii=False, indent=2), encoding="utf-8")
        question = "合成案件の中間値と最終値のMacro F1改善幅を小数第6位まで答えてください。"
        if case_id == "negative_operation_mismatch":
            question = "合成案件の中間値に対する最終値のMacro F1比率を小数第6位まで答えてください。"
        rows.append({"case_id": case_id, "expected_allowed": positive, "question": question, "expected_answer": expected, "note": note})
    write_csv(base / "cases.csv", rows)
    return rows, base


def generate_silver(root: Path) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    raw = root / "data/raw/share/share/共有ドライブ/プロジェクト"
    question_rows: list[dict[str, Any]] = []
    evidence_rows: list[dict[str, Any]] = []
    for metrics_path in sorted(raw.rglob("metrics.json")):
        if "analysis_outputs" not in metrics_path.as_posix():
            continue
        project = metrics_path.relative_to(raw).parts[0]
        metrics = json.loads(metrics_path.read_text(encoding="utf-8-sig"))
        final_value = metrics.get("accuracy")
        if not isinstance(final_value, (int, float)):
            continue
        reports = sorted((raw / project).glob("05.会議/報告資料/*.docx"))
        candidates: list[tuple[Path, Decimal]] = []
        for report in reports:
            text = docx_text(report)
            if "中間報告" not in text:
                continue
            value = metric_value(text, "accuracy", "中間")
            if value is not None:
                candidates.append((report, value))
        unique = {(str(value), report.name) for report, value in candidates}
        if not unique:
            continue
        # 同じ中間値が複数文書にある場合は、報告資料フォルダ内の最終日付を正本候補とする。
        values = {value for _, value in candidates}
        if len(values) != 1:
            continue
        report, interim = sorted(candidates, key=lambda item: item[0].name)[-1]
        final = Decimal(str(final_value)); answer = (final - interim).quantize(Decimal("0.000001"), rounding=ROUND_HALF_UP)
        silver_id = f"silver_accuracy_difference_{len(question_rows) + 1:02d}"
        question_rows.append({"silver_id": silver_id, "question": f"{project}の中間報告資料に記載されたAccuracyと最終分析出力metrics.jsonのAccuracyを用いて、改善幅を小数第6位まで答えてください。"})
        evidence_rows.append({"silver_id": silver_id, "answer": format(answer, "f"), "project": project, "interim_value": str(interim), "final_value": str(final), "interim_source": report.relative_to(root).as_posix(), "final_source": metrics_path.relative_to(root).as_posix(), "formula": "final - interim"})
        if len(question_rows) >= 5:
            break
    base = root / "evaluation/silver/calculation/cross_source_metric_difference"
    write_csv(base / "questions.csv", question_rows)
    write_csv(base / "answers.csv", [{"silver_id": row["silver_id"], "answer": row["answer"]} for row in evidence_rows])
    write_jsonl(base / "evidence.jsonl", evidence_rows)
    return question_rows, evidence_rows


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--root", type=Path, default=Path.cwd())
    parser.add_argument("--baseline-run", default="calculation_capability_baseline_fresh_v1")
    parser.add_argument("--test-run", default="semantic_phase_test_shadow_v2")
    parser.add_argument("--matrix-run", default="capability_matrix_130_v1")
    parser.add_argument("--output-run", default="calculation_capability_phase_v1")
    args = parser.parse_args(); root = args.root.resolve()
    analysis_dir = root / "data/output" / args.output_run / "analysis"; analysis_dir.mkdir(parents=True, exist_ok=True)
    matrix = read_csv(root / "data/output" / args.matrix_run / "analysis/capability_matrix_all_130.csv")
    targets = [row for row in matrix if row["primary_question_type"] == "calculation"]
    valid_eval = {int(row["question_id"]): row for row in read_csv(root / "data/output" / args.baseline_run / "evaluation/valid_evaluation.csv")}
    answers_by_dataset = {
        "valid": {int(row["question_id"]): row for row in read_jsonl(root / "data/output" / args.baseline_run / "answer_results.jsonl")},
        "test": {int(row["question_id"]): row for row in read_jsonl(root / "data/output" / args.test_run / "answer_results.jsonl")},
    }
    plans_by_dataset = {
        "valid": {int(row["question_id"]): row for row in read_jsonl(root / "data/work" / args.baseline_run / "planning/final_source_plans.jsonl")},
        "test": {int(row["question_id"]): row for row in read_jsonl(root / "data/work" / args.test_run / "planning/final_source_plans.jsonl")},
    }
    audit: list[dict[str, Any]] = []
    for row in targets:
        dataset = row["dataset"]; qid = int(row["question_id"]); question = row["question_original"]
        primary, secondary = classify_pattern(question); columns, conditions, units, rounding = required_fields(question)
        answer = answers_by_dataset[dataset].get(qid, {}); plan = plans_by_dataset[dataset].get(qid, {}); cardinality, relation = source_shape(question)
        operations = [item.get("operation_type") or item.get("tool_name") for item in plan.get("operations", [])]
        current_columns = [column for evidence in answer.get("evidence_locations", []) for column in evidence.get("input_columns", evidence.get("columns_used", []))]
        audit.append({
            "dataset": dataset, "question_id": qid, "question_original": question,
            "primary_calculation_pattern": primary, "secondary_patterns": " | ".join(secondary),
            "required_operations": " | ".join(operations), "operation_order": " -> ".join(operations),
            "source_cardinality": cardinality, "source_relation": relation,
            "required_file_roles": row["required_document_roles"], "required_file_types": row["required_file_types"],
            "required_columns": " | ".join(columns), "required_conditions": " | ".join(conditions),
            "required_join_keys": "project | metric_name | source_phase" if "cross_file_join" in secondary else "",
            "required_units": " | ".join(units), "required_rounding": " | ".join(rounding),
            "current_executor": row["current_executor"], "current_failure_stage": row["failure_stage"],
            "current_selected_files": " | ".join(answer.get("selected_files", [])), "current_selected_columns": " | ".join(current_columns),
            "deterministic_possible": primary not in {"not_actually_calculation", "unknown_calculation", "chart_source_calculation"},
            "semantic_help_required": primary in {"not_actually_calculation", "unknown_calculation"},
            "vision_required": primary == "chart_source_calculation", "valid_answer_available": dataset == "valid",
            "current_valid_result": valid_eval.get(qid, {}).get("normalized_match", "not_applicable") if dataset == "valid" else "not_applicable",
            "synthetic_testable": primary != "not_actually_calculation", "silver_testable": primary not in {"not_actually_calculation", "chart_source_calculation"},
            "shadow_gold_required": dataset == "test" and primary in {"not_actually_calculation", "unknown_calculation", "chart_source_calculation"},
            "implementation_difficulty": 3 if primary in {"difference_calculation", "multi_condition_aggregation", "id_count_or_nunique"} else 4,
            "error_risk": 2 if primary in {"difference_calculation", "single_table_aggregation"} else 3,
            "reusability": 5 if primary in {"difference_calculation", "ratio_or_percentage", "id_count_or_nunique"} else 4,
            "test_frequency": 0, "priority_score": 0.0,
        })
    counts = Counter(row["primary_calculation_pattern"] for row in audit if row["dataset"] == "test")
    valid_unresolved = Counter(row["primary_calculation_pattern"] for row in audit if row["dataset"] == "valid" and str(row["current_valid_result"]).lower() != "true")
    for row in audit:
        pattern = row["primary_calculation_pattern"]; row["test_frequency"] = counts[pattern]
        measurability = 1.0 if valid_unresolved[pattern] else 0.35
        row["priority_score"] = round((counts[pattern] + 1) * measurability * row["reusability"] * (1 if row["synthetic_testable"] else 0.3) * (1 if row["silver_testable"] else 0.4) / (row["implementation_difficulty"] * row["error_risk"]), 4)
    write_csv(analysis_dir / "calculation_capability_23.csv", audit)
    summary = []
    for pattern in CALCULATION_PATTERNS:
        items = [row for row in audit if row["primary_calculation_pattern"] == pattern]
        if items:
            summary.append({"pattern": pattern, "valid_count": sum(row["dataset"] == "valid" for row in items), "test_count": sum(row["dataset"] == "test" for row in items), "valid_unresolved": valid_unresolved[pattern], "max_priority_score": max(row["priority_score"] for row in items)})
    write_csv(analysis_dir / "calculation_pattern_summary.csv", summary)
    priority = sorted(summary, key=lambda row: float(row["max_priority_score"]), reverse=True)
    for index, row in enumerate(priority, 1): row["recommended_order"] = index
    write_csv(analysis_dir / "calculation_vertical_slice_priority.csv", priority)
    selected = priority[0]
    (analysis_dir / "selected_calculation_slice.md").write_text(
        "# Selected Calculation Slice\n\n"
        f"- 選定した操作パターン: {selected['pattern']}\n- 対象valid: {selected['valid_count']}問\n- 対象test: {selected['test_count']}問\n"
        "- 現在の失敗原因: 複数資料をsingle sourceとして計画し、単一値抽出で停止している。\n"
        "- 必要な汎用機能: 情報源役割分離、同一指標の値抽出、source phase順序、差分、Decimal丸め、独立再計算。\n"
        "- 期待valid増分: 1問\n- 期待testカバレッジ: 同一指標を複数資料から取得する改善幅質問。現行testの1問は複合差分・除算を含むため別仕様として抑制する。\n"
        "- 実装リスク: 同一指標の複数値、改善方向、単位不一致。\n"
        "- 評価方法: Synthetic正負例、raw由来Silver、valid回帰、test Shadow Audit。\n",
        encoding="utf-8",
    )
    audit_lines = ["# Calculation Classification Audit", "", "## 凡例", "", "patternは質問が要求する主たる計算操作、secondaryは追加で必要な操作を表します。", "", "## 修正点", "", "- validの残余リスク抽出は計算ではなくsemantic status lookupへ再分類。", "- 改善幅質問はsingle sourceではなくcross-file joinを必要とする。", "- 発行数・タスクID数はsumではなくcountまたはnuniqueとして扱う。", "- 係数予測はrow lookupへフォールバックしてはならない。", ""]
    (analysis_dir / "calculation_classification_audit.md").write_text("\n".join(audit_lines), encoding="utf-8")

    synthetic_rows, _ = generate_synthetic(root)
    silver_questions, silver_evidence = generate_silver(root)
    write_csv(analysis_dir / "silver_calculation_questions.csv", silver_questions)
    shadow = [row for row in audit if row["dataset"] == "test" and row["primary_calculation_pattern"] == selected["pattern"]][:3]
    shadow_rows = [{"question_id": row["question_id"], "question": row["question_original"], "selection_reason": "selected slice representative", "required_files": row["required_file_roles"], "required_operation": row["primary_calculation_pattern"], "independent_verification_possible": row["deterministic_possible"], "shadow_gold_difficulty": row["implementation_difficulty"]} for row in shadow]
    write_csv(analysis_dir / "calculation_shadow_gold_candidates.csv", shadow_rows)

    baseline_eval = read_csv(root / "data/output" / args.baseline_run / "evaluation/valid_evaluation.csv")
    write_csv(analysis_dir / "baseline_16_results.csv", baseline_eval)
    baseline_answers = answers_by_dataset["valid"]
    routes = []
    for row in baseline_eval:
        if str(row.get("normalized_match", "")).lower() != "true": continue
        answer = baseline_answers[int(row["question_id"])]
        routes.append({"question_id": row["question_id"], "executor": " | ".join(answer.get("operations_executed", [])), "actual_used_files": " | ".join(answer.get("selected_files", [])), "evidence": json.dumps(answer.get("evidence_locations", []), ensure_ascii=False), "gate_status": answer.get("gate_status", ""), "final_answer": answer.get("answer", "")})
    write_csv(analysis_dir / "baseline_16_routes.csv", routes)
    manifest = json.loads((root / "data/output" / args.baseline_run / "run_manifest.json").read_text(encoding="utf-8"))
    (analysis_dir / "baseline_16_quality.md").write_text(f"# Baseline 16 Quality\n\n- correct: 16\n- incorrect: 0\n- blank: 14\n- score: +16\n- raw files: {manifest.get('raw_file_count')}\n- cache hit / miss: {manifest.get('cache_hits')} / {manifest.get('cache_misses')}\n", encoding="utf-8")
    git = ["git", "-c", "safe.directory=E:/PC/デスクトップ/SIGNATE/SIGNATE_Agentic_RAG"]
    for name, command in (("pre_calculation_git_status.txt", git + ["status", "--short"]), ("pre_calculation_diff.patch", git + ["diff"]), ("pre_calculation_diff_stat.txt", git + ["diff", "--stat"]), ("pre_calculation_changed_files.txt", git + ["diff", "--name-only"])):
        result = subprocess.run(command, cwd=root, capture_output=True, text=True, encoding="utf-8", errors="replace")
        (analysis_dir / name).write_text(result.stdout + result.stderr, encoding="utf-8")
    print(json.dumps({"calculation_questions": len(audit), "selected": selected["pattern"], "synthetic_cases": len(synthetic_rows), "silver_questions": len(silver_questions), "shadow_candidates": len(shadow_rows)}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
