from __future__ import annotations

import json
import shutil
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path


try:
    import msoffcrypto  # noqa: F401
    import openpyxl
    from docx import Document
    from msoffcrypto.format.ooxml import OOXMLFile

    HAS_OFFICE_DEPS = True
except Exception:
    HAS_OFFICE_DEPS = False


def encrypt_ooxml(plain_path: Path, encrypted_path: Path, password: str) -> None:
    with plain_path.open("rb") as src, encrypted_path.open("wb") as dst:
        OOXMLFile(src).encrypt(password, dst)


def make_docx(path: Path, lines: list[str]) -> None:
    document = Document()
    for line in lines:
        document.add_paragraph(line)
    document.save(path)


def make_xlsx(path: Path, rows: list[list[object]]) -> None:
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    for row in rows:
        sheet.append(row)
    workbook.save(path)


class ProtectedOfficePipelineTest(unittest.TestCase):
    @unittest.skipUnless(HAS_OFFICE_DEPS, "Office encryption dependencies are not installed")
    def test_protected_office_files_are_resolved_from_raw_only(self) -> None:
        root = Path(__file__).resolve().parents[1]
        run_id = "unittest_protected_source_selection"
        work_dir = root / "data" / "work" / run_id
        output_dir = root / "data" / "output" / run_id
        shutil.rmtree(work_dir, ignore_errors=True)
        shutil.rmtree(output_dir, ignore_errors=True)

        secrets = ["NameSecret123", "PassSecret123", "WideSecret123", "DA-TEST-20250806-xlsx"]
        with tempfile.TemporaryDirectory() as tmp:
            tmp_root = Path(tmp)
            raw_root = tmp_root / "共有ドライブ"
            management_dir = raw_root / "社内管理"
            project_root = raw_root / "プロジェクト" / "テスト株式会社"
            contract_dir = project_root / "01.契約"
            plan_dir = project_root / "02.計画"
            management_dir.mkdir(parents=True)
            contract_dir.mkdir(parents=True)
            plan_dir.mkdir(parents=True)

            glossary = Document()
            table = glossary.add_table(rows=1, cols=4)
            table.rows[0].cells[0].text = "案件名"
            table.rows[0].cells[1].text = "主略称"
            table.rows[0].cells[2].text = "別名候補"
            table.rows[0].cells[3].text = "補足"
            row = table.add_row().cells
            row[0].text = "テスト株式会社"
            row[1].text = "TEST"
            row[2].text = "テスト"
            row[3].text = "TESTを正式"
            glossary.save(management_dir / "社内用語集.docx")
            make_docx(
                management_dir / "データアステル社内規定_パスワード導出規則.docx",
                [
                    "DA-[案件略号]-[開始年月日8桁]-[拡張子コード]",
                    "社内用語集にて規定されている主略称を使用する",
                    "契約開始日のYYYYMMDDを使う",
                ],
            )

            plain_contract = contract_dir / "contract_plain.docx"
            make_docx(
                plain_contract,
                [
                    "契約開始日: 2025年8月5日",
                    "開始日: 2025年8月6日",
                    "この契約書から開始日候補を取得する。",
                ],
            )
            encrypted_contract = contract_dir / "契約書_pw=NameSecret123.docx"
            encrypt_ooxml(plain_contract, encrypted_contract, "NameSecret123")
            plain_contract.unlink()

            for filename, password, value in [
                ("資料_password=PassSecret123.xlsx", "PassSecret123", 10),
                ("資料_ＰＷ=WideSecret123.xlsx", "WideSecret123", 20),
                ("スケジュール.xlsx", "DA-TEST-20250806-xlsx", 30),
                ("見積_pw=wrong.xlsx", "DA-TEST-20250806-xlsx", 40),
            ]:
                plain = plan_dir / f"{filename}.plain.xlsx"
                make_xlsx(plain, [["項目", "値"], ["A", value]])
                encrypted = plan_dir / filename
                encrypt_ooxml(plain, encrypted, password)
                plain.unlink()
            (plan_dir / "~$一時.xlsx").write_bytes(b"temporary")
            raw_hashes = {path: path.read_bytes() for path in raw_root.rglob("*") if path.is_file()}

            questions_path = tmp_root / "questions_valid.csv"
            questions_path.write_text("index,question,answer\n0,テスト株式会社の資料を確認してください,ok\n", encoding="utf-8")

            result = subprocess.run(
                [
                    sys.executable,
                    "-m",
                    "rag_competition.pipeline",
                    "--mode",
                    "source-selection",
                    "--fresh",
                    "--split",
                    "valid",
                    "--run-id",
                    run_id,
                    "--raw-root",
                    str(raw_root),
                    "--questions-path",
                    str(questions_path),
                    "--api-mode",
                    "off",
                    "--resolve-protected-files",
                    "--no-render-pdf-pages",
                ],
                cwd=root,
                text=True,
                capture_output=True,
                check=False,
            )
            self.assertEqual(result.returncode, 0, result.stderr + result.stdout)
            for path, before in raw_hashes.items():
                self.assertEqual(path.read_bytes(), before)

        manifest = json.loads((work_dir / "run_manifest.json").read_text(encoding="utf-8"))
        self.assertEqual(manifest["status"], "completed")
        self.assertGreaterEqual(manifest["protected_file_count"], 5)
        self.assertGreaterEqual(manifest["filename_password_success_count"], 3)
        self.assertGreaterEqual(manifest["rule_derived_password_success_count"], 2)
        self.assertEqual(manifest["temporary_office_file_count"], 1)

        protected_dir = work_dir / "protected_files"
        protected_text = "\n".join(
            (protected_dir / name).read_text(encoding="utf-8")
            for name in ["protected_file_inventory.jsonl", "password_candidates.jsonl", "decryption_attempts.jsonl", "decryption_summary.csv"]
        )
        for secret in secrets + ["wrong"]:
            self.assertNotIn(secret, protected_text)
        self.assertIn("***", protected_text)

        decrypted_files = list((work_dir / "decrypted").rglob("*.*"))
        self.assertGreaterEqual(len(decrypted_files), 5)
        self.assertTrue(all(work_dir in path.parents for path in decrypted_files))

        attempts = [json.loads(line) for line in (protected_dir / "decryption_attempts.jsonl").read_text(encoding="utf-8").splitlines() if line]
        self.assertTrue(any(not row["success"] for row in attempts))
        self.assertTrue(any(row["success"] and row["validation_success"] for row in attempts))

        records = [json.loads(line) for line in (protected_dir / "protected_file_inventory.jsonl").read_text(encoding="utf-8").splitlines() if line]
        self.assertTrue(any(row["resolution_status"] == "temporary_file" for row in records))
        self.assertTrue(any(row["resolution_status"] == "resolved_from_rule" for row in records))
        self.assertTrue(all(row["extraction_success"] or row["resolution_status"] == "temporary_file" for row in records if row["requires_password"] or row["is_temporary_file"]))

    @unittest.skipUnless(HAS_OFFICE_DEPS, "Office encryption dependencies are not installed")
    def test_ambiguous_project_alias_is_not_decrypted(self) -> None:
        from rag_competition.inventory import build_inventory
        from rag_competition.protected_files import resolve_protected_files

        root = Path(__file__).resolve().parents[1]
        run_id = "unittest_protected_ambiguous"
        run_dir = root / "data" / "work" / run_id
        shutil.rmtree(run_dir, ignore_errors=True)

        with tempfile.TemporaryDirectory() as tmp:
            tmp_root = Path(tmp)
            raw_root = tmp_root / "共有ドライブ"
            management_dir = raw_root / "社内管理"
            project_dir = raw_root / "プロジェクト" / "曖昧株式会社" / "02.計画"
            management_dir.mkdir(parents=True)
            project_dir.mkdir(parents=True)

            glossary = Document()
            table = glossary.add_table(rows=1, cols=4)
            for cell, value in zip(table.rows[0].cells, ["案件名", "主略称", "別名候補", "補足"]):
                cell.text = value
            for alias in ["AMB1", "AMB2"]:
                row = table.add_row().cells
                row[0].text = "曖昧株式会社"
                row[1].text = alias
                row[2].text = "曖昧"
                row[3].text = "同一案件に複数主略称があるテスト"
            glossary.save(management_dir / "社内用語集.docx")
            make_docx(
                management_dir / "データアステル社内規定_パスワード導出規則.docx",
                [
                    "DA-[案件略号]-[開始年月日8桁]-[拡張子コード]",
                    "社内用語集にて規定されている主略称を使用する",
                    "契約開始日のYYYYMMDDを使う",
                ],
            )
            plain = project_dir / "plain.xlsx"
            make_xlsx(plain, [["契約開始日", "2025年8月6日"]])
            encrypted = project_dir / "スケジュール.xlsx"
            encrypt_ooxml(plain, encrypted, "DA-AMB1-20250806-xlsx")
            plain.unlink()

            files = build_inventory(raw_root, root, run_dir / "inventory")
            result = resolve_protected_files(files, root, raw_root, run_dir)

        records = [record for record in result.records if record.requires_password]
        self.assertEqual(len(records), 1)
        self.assertEqual(records[0].resolution_status, "ambiguous_alias")
        self.assertFalse(records[0].decryption_success)
        self.assertEqual(len(result.attempts), 0)


if __name__ == "__main__":
    unittest.main()
