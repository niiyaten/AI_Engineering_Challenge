from docx import Document

from rag_recovery.executors.direct import DirectStructuredExecutor
from rag_recovery.executors.remaining50_generalization import Remaining50GeneralizationExecutor
from rag_recovery.models import QueryPlan, Question
from rag_recovery.store import DocumentStore


def test_mortality_ratio_uses_displayed_rank_not_numeric_sort(sample_share):
    path = sample_share / "共有ドライブ" / "プロジェクト" / "株式会社テスト分析" / "00.提案" / "糖尿病統計情報.docx"
    doc = Document()
    table = doc.add_table(rows=6, cols=5)
    rows = [
        ["順位", "死亡率が高い都道府県（ワースト）", "死亡率（%）", "死亡率が低い都道府県（ベスト）", "死亡率（%）"],
        ["1位", "青森県", "18.2", "神奈川県", "7.2"],
        ["2位", "秋田県", "16.3", "愛知県", "7.9"],
        ["3位", "香川県", "16.1", "東京都", "8.8"],
        ["4位", "鹿児島県", "15.0", "滋賀県", "7.3"],
        ["5位", "徳島県", "14.9", "奈良県", "-"],
    ]
    for r, values in zip(table.rows, rows):
        for c, value in zip(r.cells, values):
            c.text = value
    doc.save(path)

    store = DocumentStore(sample_share)
    q = Question("test", 99, "テスト分析の糖尿病統計情報調査結果において、死亡率が最も高い都道府県の死亡率は、4番目に低い都道府県の死亡率の何倍ですか。小数第2位まで求めてください。")
    result = Remaining50GeneralizationExecutor().execute(q, QueryPlan("remaining50_generalization", project_hints=("テスト分析",)), store)

    assert result.answered
    assert result.answer == "2.49倍"
    assert result.method == "displayed_rank_mortality_ratio"


def test_contract_bold_fields_use_explicit_separator(sample_share):
    path = sample_share / "共有ドライブ" / "プロジェクト" / "株式会社テスト分析" / "01.契約" / "契約書.docx"
    doc = Document()
    for value in [
        "2025年10月1日",
        "time_and_materials",
        "実績工数に基づき、案件完了後に最終成果物の検収を経て一括精算する。",
        "30分単位",
        "25,000円／時間",
    ]:
        run = doc.add_paragraph().add_run(value)
        run.bold = True
    doc.save(path)

    store = DocumentStore(sample_share)
    q = Question("test", 3, "テスト分析の契約書において、太字で記載されている箇所のうち、日付以外のものをすべて抽出してください。")
    result = DirectStructuredExecutor().execute(q, QueryPlan("direct", project_hints=("テスト分析",)), store)

    assert result.answered
    assert result.answer == "time_and_materials／実績工数に基づき、案件完了後に最終成果物の検収を経て一括精算する。／30分単位／25,000円／時間"
    assert "。、" not in result.answer
